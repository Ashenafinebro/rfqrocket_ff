import os
import json
import time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import openai
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

load_dotenv()

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['PROCESSED_FOLDER'] = 'processed'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024
app.config['SMTP_SERVER'] = os.getenv('SMTP_SERVER')
app.config['SMTP_PORT'] = int(os.getenv('SMTP_PORT', 587))
app.config['SMTP_USERNAME'] = os.getenv('SMTP_USERNAME')
app.config['SMTP_PASSWORD'] = os.getenv('SMTP_PASSWORD')
app.config['EMAIL_FROM'] = os.getenv('EMAIL_FROM')
app.config['THREAD_POOL_WORKERS'] = 4

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)

openai.api_key = os.getenv('OPENAI_API_KEY')
executor = ThreadPoolExecutor(max_workers=app.config['THREAD_POOL_WORKERS'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_file(filepath, filename):
    text = ""
    try:
        if filename.endswith('.pdf'):
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif filename.endswith('.docx'):
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text])
        elif filename.endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        return text.strip()
    except Exception as e:
        print(f"Error extracting text: {str(e)}")
        return ""

def split_text(text, max_length=5000):
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_length = len(para)
        if current_length + para_length > max_length:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_length = para_length
        else:
            current_chunk.append(para)
            current_length += para_length
    
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks

def process_chunk(chunk):
    prompt = f"""Analyze this government solicitation document chunk and extract all possible information to create a comprehensive Request for Quotation (RFQ). Extract all sections, requirements, specifications, terms, and any relevant details. Structure the output as follows:

1. GENERAL INFORMATION: Include solicitation number, title, agency, date, etc.
2. REQUIREMENTS: Detailed technical requirements and specifications
3. DELIVERABLES: List of all required deliverables
4. PERIOD_OF_PERFORMANCE: Start/end dates or duration
5. EVALUATION_CRITERIA: How proposals will be evaluated
6. SUBMISSION_REQUIREMENTS: Format, deadlines, submission instructions
7. TERMS_AND_CONDITIONS: Contractual terms and conditions
8. CONTACT_INFORMATION: Points of contact

Document Text:
{chunk}

Provide the output in JSON format with these exact top-level keys. Include all relevant details under each section."""
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-1106-preview",
            messages=[
                {"role": "system", "content": "You are an AI assistant generating vendor-safe RFQs from government solicitations. Extract only the info vendors need to quote—scope, specs, quantities, delivery, performance period, city/state location, evaluation criteria, and submission format. Remove all agency names, contacts, solicitation numbers, compliance references (FAR, SAM), and base-identifying details. Rewrite the output as a clean, private RFQ with no government identifiers"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"Error processing chunk: {str(e)}")
        return {}

def generate_rfq(text):
    start_time = time.time()
    text_chunks = split_text(text)
    rfq_data = {
        'GENERAL_INFORMATION': {},
        'REQUIREMENTS': [],
        'DELIVERABLES': [],
        'PERIOD_OF_PERFORMANCE': {},
        'EVALUATION_CRITERIA': [],
        'SUBMISSION_REQUIREMENTS': [],
        'TERMS_AND_CONDITIONS': [],
        'CONTACT_INFORMATION': {}
    }
    
    futures = [executor.submit(process_chunk, chunk) for chunk in text_chunks]
    for future in futures:
        chunk_data = future.result()
        for section in rfq_data.keys():
            if section in chunk_data:
                if isinstance(rfq_data[section], dict):
                    rfq_data[section].update(chunk_data[section])
                elif isinstance(rfq_data[section], list):
                    if isinstance(chunk_data[section], list):
                        rfq_data[section].extend(chunk_data[section])
                    else:
                        rfq_data[section].append(chunk_data[section])
    
    print(f"RFQ generation completed in {time.time() - start_time:.2f} seconds")
    return rfq_data

def set_doc_styles(doc):
    styles = doc.styles
    font_name = 'Calibri'
    
    for style in styles:
        if style.type == 1:
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 2'].font.bold = True
    doc.styles['Heading 2'].font.color.rgb = RGBColor(0, 0, 0)

def add_custom_header(doc, text):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    doc.add_paragraph()

def add_section(doc, title, content, level=1):
    doc.add_heading(title, level=level)
    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                for key, value in item.items():
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value))
            else:
                doc.add_paragraph(str(item), style='List Bullet')
    elif isinstance(content, dict):
        for key, value in content.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(str(value))
    else:
        doc.add_paragraph(str(content))

def create_rfq_document(rfq_data, output_path):
    try:
        doc = Document()
        set_doc_styles(doc)
        add_custom_header(doc, "RFQ Generated by RFQRocket")
        
        doc.add_heading('REQUEST FOR QUOTATION', level=0)
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph()
        
        add_table_of_contents(doc)
        
        section_order = [
            ('GENERAL_INFORMATION', '1. General Information'),
            ('REQUIREMENTS', '2. Technical Requirements'),
            ('DELIVERABLES', '3. Deliverables'),
            ('PERIOD_OF_PERFORMANCE', '4. Period of Performance'),
            ('EVALUATION_CRITERIA', '5. Evaluation Criteria'),
            ('SUBMISSION_REQUIREMENTS', '6. Submission Requirements'),
            ('TERMS_AND_CONDITIONS', '7. Terms and Conditions'),
            ('CONTACT_INFORMATION', '8. Contact Information')
        ]
        
        for section_key, section_title in section_order:
            if section_key in rfq_data and rfq_data[section_key]:
                add_section(doc, section_title, rfq_data[section_key], level=1)
        
        doc.add_page_break()
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating document: {str(e)}")
        return False

def send_email_with_attachment(to_email, subject, body, attachment_path):
    try:
        msg = MIMEMultipart()
        msg['From'] = app.config['EMAIL_FROM']
        msg['To'] = to_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain'))
        
        with open(attachment_path, 'rb') as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment_path)}')
            msg.attach(part)
        
        with smtplib.SMTP(app.config['SMTP_SERVER'], app.config['SMTP_PORT']) as server:
            server.starttls()
            server.login(app.config['SMTP_USERNAME'], app.config['SMTP_PASSWORD'])
            server.send_message(msg)
        
        return True
    except Exception as e:
        print(f"Error sending email: {str(e)}")
        return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload')
def upload_page():
    return render_template('upload.html')

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(upload_path)
        
        try:
            text = extract_text_from_file(upload_path, filename)
            if not text:
                return jsonify({'error': 'Could not extract text from file'}), 500
            
            rfq_data = generate_rfq(text)
            if not rfq_data:
                return jsonify({'error': 'Failed to generate RFQ data'}), 500
            
            output_filename = f"RFQ_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
            
            if not create_rfq_document(rfq_data, output_path):
                return jsonify({'error': 'Failed to create RFQ document'}), 500
            
            return jsonify({
                'success': True,
                'download_url': f'/download/{output_filename}',
                'filename': output_filename
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        finally:
            if os.path.exists(upload_path):
                os.remove(upload_path)
    else:
        return jsonify({'error': 'File type not allowed'}), 400

@app.route('/api/send-email', methods=['POST'])
def send_email():
    data = request.get_json()
    if not data or 'email' not in data or 'filename' not in data:
        return jsonify({'error': 'Missing required fields'}), 400
    
    email = data['email']
    filename = data['filename']
    filepath = os.path.join(app.config['PROCESSED_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    subject = "Your RFQ Document from RFQRocket"
    body = f"""Dear Recipient,

Please find attached the Request for Quotation (RFQ) document generated from your source file.

Document: {filename}
Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}

If you have any questions or need further assistance, please don't hesitate to contact us.

Best regards,
The RFQRocket Team"""
    
    if send_email_with_attachment(email, subject, body, filepath):
        return jsonify({'success': True})
    else:
        return jsonify({'error': 'Failed to send email'}), 500

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['PROCESSED_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    app.run(host='0.0.0.0', port=port)